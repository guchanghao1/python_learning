# !/usr/bin/env python
# -*- coding: utf-8 -*-
''''''
# ----------------------------------------------------------------------------------------------------------------------
from docx import Document  # 一个word文件就是一个文档
from docx.shared import Pt, RGBColor  # 字体大小
from docx.enum.text import WD_ALIGN_PARAGRAPH  # 段落对齐方式
from docx.oxml.ns import qn  # 跟字体有关的模块

# 批量处理
info = [
    {'name': '张三', 'send_type': '电子邮件', 'offer': '企划'},
    {'name': '李四', 'send_type': '学校推荐', 'offer': '信息管理'},
    {'name': '王五', 'send_type': '电话沟通', 'offer': '企划'},
    {'name': '陈六', 'send_type': '电子邮件', 'offer': '信息管理'},
    {'name': '马奇', 'send_type': '网站表格', 'offer': '广告推广'},
]


# 批量操作
def creat_parag(para, content):
    para.paragraph_format.first_line_indent = Pt(32)  # 段落的首行缩进
    run = para.add_run(content)
    run.font.name = u'仿宋'
    run0.element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    run0.font.size = Pt(16)
    run0.font.bol = True


for person in info:
    vdoc = Document()  # 创建对象
    # 创建一个段落
    p0 = vdoc.add_paragraph()
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中对齐
    p0_format = p0.paragraph_format  # 段落格式的属性对象
    p0_format.space_before = Pt(8)
    p0_format.space_after = Pt(32)
    run0 = p0.add_run('面试通知书')
    run0.font.name = u'黑体'
    run0.element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    run0.font.size = Pt(21)
    run0.font.bol = True
    # 第一个段落
    p1 = vdoc.add_paragraph()
    run1 = p1.add_run(f'{person.get("name")}先生/女士：你好')
    run1.font.name = u'仿宋'
    run0.element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    run0.font.size = Pt(16)
    run0.font.bol = True

    p2 = vdoc.add_paragraph()
    p2.paragraph_format.first_line_indent = Pt(32)  # 段落的首行缩进
    run2 = p2.add_run(f'{person.get("send_type")}--{person.get("offer")}')
    run2.font.name = u'仿宋'
    run0.element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    run0.font.size = Pt(16)
    run0.font.bol = True

    p3 = vdoc.add_paragraph()
    creat_parag(p3, '1')
    p4 = vdoc.add_paragraph()
    creat_parag(p4, '2')
    p5 = vdoc.add_paragraph()
    creat_parag(p5, '3')
    p6 = vdoc.add_paragraph()
    creat_parag(p6, '4')
    p7 = vdoc.add_paragraph()
    p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
    creat_parag(p7, 'XXX人力资源部')
    p8 = vdoc.add_paragraph()
    p8.alignment = WD_ALIGN_PARAGRAPH.CENTER
    creat_parag(p8, '2025.12.1')
    vdoc.save(f'./21word数据批量转换/{person.get("name")}.docx')
